"""
Resume Rewriter — AI rewrites resume sections based on analysis gaps.
Generates a clean ATS-safe DOCX and PDF for download. Pro-only feature.
"""
import io
import json
import logging
import uuid
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor

from app.ai.llm_client import llm_client
from app.ai.prompt_templates import get_prompt
from app.ai.cost_governor import cost_governor

logger = logging.getLogger(__name__)


class RewriteResult:
    def __init__(self):
        self.docx_url: str = ""
        self.pdf_url: str = ""
        self.sections_changed: list = []
        self.changes_summary: str = ""
        self.model_used: str = ""
        self.tokens_used: int = 0
        self.cost_usd: float = 0.0
        self.improved_text: str = ""   # flat text of improved resume for ATS re-scoring
        self.projected_ats_score: float = 0.0
        self.error: Optional[str] = None


class ResumeRewriter:

    async def rewrite(
        self,
        resume: dict,
        analyses: list,
        user_id: str,
    ) -> RewriteResult:
        result = RewriteResult()

        # Check budget
        budget_check = await cost_governor.can_use_llm(user_id, "pro", "rewrite")
        if not budget_check.get("allowed", True):
            result.error = "AI budget exhausted for this month"
            return result

        # Build context from existing analyses
        gaps_context = self._build_gaps_context(analyses)
        recommended_changes = self._build_recommended_changes(analyses)
        missing_keywords = self._build_missing_keywords(analyses)

        # Use raw_text if structured_data is sparse (< 3 sections)
        structured = resume.get("structured_data") or {}
        raw_text = resume.get("raw_text", "")
        if len(structured) < 3 and raw_text:
            original_sections_text = raw_text
        else:
            original_sections_text = json.dumps(structured, indent=2)

        # Build prompt
        prompt_template = get_prompt("rewrite")
        prompt = prompt_template.format(
            original_sections=original_sections_text,
            missing_keywords=missing_keywords,
            gaps_context=gaps_context,
            recommended_changes=recommended_changes,
        )

        # LLM call — 8192 tokens needed: full resume sections can be 4-6k tokens
        try:
            response = await llm_client.complete(
                prompt=prompt,
                model="gemini-2.5-flash",
                max_tokens=8192,
                temperature=0.4,
            )
        except Exception as e:
            logger.error(f"LLM rewrite API call failed: {e}")
            result.error = f"AI rewrite failed: {str(e)}"
            return result

        try:
            rewritten = response.as_json()
        except Exception as json_err:
            raw_preview = (response.text or "")[:800]
            logger.error(f"Rewrite JSON parse failed: {json_err} | Raw response: {raw_preview}")
            result.error = f"AI response parse failed: {str(json_err)}"
            return result

        result.model_used = response.model
        result.tokens_used = response.tokens
        result.cost_usd = response.cost_usd
        result.sections_changed = rewritten.get("sections_changed", [])
        result.changes_summary = rewritten.get("changes_summary", "")

        # Record spend
        try:
            await cost_governor.record_spend(user_id, response.cost_usd)
        except Exception as e:
            logger.warning(f"Cost recording failed (non-fatal): {e}")

        # Merge: start from structured data (or empty), overlay LLM rewrites
        # LLM output takes priority — it's the optimized version
        merged = {}
        # First put structured data as fallback
        if isinstance(structured, dict):
            merged.update(structured)
        # Then overlay ALL LLM rewrites (skip metadata keys)
        skip_keys = {"sections_changed", "changes_summary"}
        for key, val in rewritten.items():
            if key in skip_keys:
                continue
            if val and isinstance(val, str) and val.strip():
                merged[key] = val

        raw_text = resume.get("raw_text", "")

        # Build improved_text (flat) for projected ATS scoring
        result.improved_text = "\n".join(
            str(v) for k, v in merged.items()
            if k not in ("sections_changed", "changes_summary") and v and isinstance(v, str)
        )
        # Projected ATS: check how many previously-missing keywords now appear
        result.projected_ats_score = self._compute_projected_score(
            result.improved_text, analyses
        )

        # Build DOCX
        try:
            docx_bytes = self._build_docx(merged, raw_text)
        except Exception as e:
            logger.error(f"DOCX build failed: {e}")
            result.error = f"Document generation failed: {str(e)}"
            return result

        # Build PDF
        try:
            pdf_bytes = self._build_pdf(merged, raw_text)
        except Exception as e:
            logger.warning(f"PDF build failed (non-fatal): {e}")
            pdf_bytes = None

        # Upload to R2
        try:
            from app.api.v1.resumes import upload_to_r2
            file_id = uuid.uuid4()
            result.docx_url = await upload_to_r2(
                docx_bytes,
                f"{user_id}/improved_{file_id}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            if pdf_bytes:
                result.pdf_url = await upload_to_r2(
                    pdf_bytes,
                    f"{user_id}/improved_{file_id}.pdf",
                    "application/pdf",
                )
        except Exception as e:
            logger.error(f"R2 upload failed: {e}")
            result.error = f"File upload failed: {str(e)}"
            return result

        return result

    def _compute_projected_score(self, improved_text: str, analyses: list) -> float:
        """Estimate new ATS score by checking keyword coverage in the improved resume text."""
        from app.ai.rule_engine import rule_engine
        text_lower = improved_text.lower()

        for analysis in analyses:
            if analysis.get("analysis_type") == "ats":
                kw = analysis.get("keyword_analysis") or {}
                required_found_before = kw.get("required_found") or []
                required_missing_before = kw.get("required_missing") or []
                preferred_found_before = kw.get("preferred_found") or []
                preferred_missing_before = kw.get("preferred_missing") or []

                all_required = required_found_before + required_missing_before
                all_preferred = preferred_found_before + preferred_missing_before

                if not all_required:
                    break

                now_required = [kw for kw in all_required if kw.lower() in text_lower]
                now_preferred = [kw for kw in all_preferred if kw.lower() in text_lower]

                req_pct = len(now_required) / len(all_required)
                pref_pct = len(now_preferred) / max(len(all_preferred), 1)
                keyword_score = round((req_pct * 85) + (pref_pct * 15), 1)

                # Format score from original (unchanged), semantic score assumed 70
                original_format = analysis.get("format_score") or 75.0
                overall = round(
                    keyword_score * 0.40
                    + 70.0 * 0.25  # semantic (estimated)
                    + original_format * 0.20
                    + 70.0 * 0.15,  # experience (estimated)
                    1
                )
                return min(98.0, overall)

        return 0.0

    def _build_missing_keywords(self, analyses: list) -> str:
        for analysis in analyses:
            if analysis.get("analysis_type") == "ats":
                kw = analysis.get("keyword_analysis") or {}
                required_missing = kw.get("required_missing") or []
                preferred_missing = kw.get("preferred_missing") or []
                lines = []
                if required_missing:
                    lines.append("REQUIRED (must add): " + ", ".join(required_missing))
                if preferred_missing:
                    lines.append("PREFERRED (add if natural): " + ", ".join(preferred_missing))
                if lines:
                    return "\n".join(lines)
        return "No specific keywords identified — improve overall quality."

    def _build_gaps_context(self, analyses: list) -> str:
        lines = []
        for analysis in analyses:
            atype = analysis.get("analysis_type", "")
            if atype == "ats":
                gaps = analysis.get("gaps") or []
                improvements = analysis.get("improvements") or []
                quick_fixes = analysis.get("quick_fixes") or []
                if gaps:
                    lines.append("ATS GAPS:")
                    for g in gaps[:8]:
                        sev = g.get("severity", "")
                        desc = g.get("description", g.get("keyword", ""))
                        fix = g.get("fix", "")
                        lines.append(f"  [{sev.upper()}] {desc}" + (f" → Fix: {fix}" if fix else ""))
                if quick_fixes:
                    lines.append("QUICK FIXES:")
                    for qf in quick_fixes[:5]:
                        lines.append(f"  • {qf.get('action', '')} (impact: {qf.get('impact', '')})")
                if improvements:
                    lines.append("IMPROVEMENTS:")
                    for imp in improvements[:5]:
                        lines.append(f"  • {imp.get('action', '')} in {imp.get('section', 'resume')}")
        return "\n".join(lines) if lines else "No specific gaps identified — general quality improvement."

    def _build_recommended_changes(self, analyses: list) -> str:
        lines = []
        for analysis in analyses:
            if analysis.get("analysis_type") == "recruiter":
                signals = analysis.get("recruiter_signals") or {}
                changes = signals.get("recommended_changes") or signals.get("recommended_resume_changes") or []
                for c in changes[:5]:
                    before = c.get("before_example", "")
                    after = c.get("after_example", "")
                    change = c.get("change", "")
                    if before and after:
                        lines.append(f"• {change}\n  BEFORE: {before}\n  AFTER: {after}")
                    elif change:
                        lines.append(f"• {change}")
        return "\n".join(lines) if lines else "No specific before/after examples available."

    def _build_docx(self, sections: dict, raw_text: str) -> bytes:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

        VIOLET = RGBColor(0x5B, 0x4F, 0xCF)
        BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
        GRAY   = RGBColor(0x55, 0x55, 0x55)
        FONT   = "Calibri"

        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Pt(54)
            sec.bottom_margin = Pt(54)
            sec.left_margin   = Pt(64)
            sec.right_margin  = Pt(64)

        def _add_border_bottom(p, color_hex="5B4FCF"):
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), color_hex)
            pBdr.append(bottom)
            pPr.append(pBdr)

        def _add_right_tab(p):
            """Add a right-aligned tab stop at the page margin."""
            pPr = p._p.get_or_add_pPr()
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "right")
            tab.set(qn("w:pos"), "9360")   # ~6.5 inch content width in twips
            tabs.append(tab)
            pPr.append(tabs)

        def heading(title):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = VIOLET
            run.font.name = FONT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            _add_border_bottom(p)
            return p

        def name_line(text):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = BLACK
            run.font.name = FONT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def contact_line(text):
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
            run.font.name = FONT
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)

        def job_title_line(title, company, date_range):
            """Role • Company (italic gray) [TAB] Date"""
            p = doc.add_paragraph()
            _add_right_tab(p)
            r1 = p.add_run(title)
            r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = BLACK; r1.font.name = FONT
            r2 = p.add_run("  •  ")
            r2.font.size = Pt(10); r2.font.color.rgb = GRAY; r2.font.name = FONT
            r3 = p.add_run(company)
            r3.italic = True; r3.font.size = Pt(10); r3.font.color.rgb = GRAY; r3.font.name = FONT
            r4 = p.add_run("\t" + date_range)
            r4.font.size = Pt(10); r4.font.color.rgb = GRAY; r4.font.name = FONT
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after  = Pt(2)

        def bullet_line(text):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = BLACK
            run.font.name = FONT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Pt(12)

        def body_text(text, size=9.5, color=None, italic=False):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.color.rgb = color or BLACK
            run.font.name = FONT
            run.italic = italic
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            return p

        def education_line(degree, institution, date_range):
            p = doc.add_paragraph()
            _add_right_tab(p)
            r1 = p.add_run(degree)
            r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = BLACK; r1.font.name = FONT
            r2 = p.add_run(", " + institution)
            r2.font.size = Pt(9.5); r2.font.color.rgb = GRAY; r2.font.name = FONT
            r3 = p.add_run("\t" + date_range)
            r3.font.size = Pt(9.5); r3.font.color.rgb = GRAY; r3.font.name = FONT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(2)

        # ── Fallback: no structured data ────────────────────────────────
        if not sections or len(sections) < 2:
            for line in raw_text.split("\n"):
                line = line.strip()
                if line:
                    body_text(line)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        # ── HEADER ──────────────────────────────────────────────────────
        hdr = sections.get("header", "")
        lines = [l.strip() for l in hdr.split("\n") if l.strip()]
        if lines:
            name_line(lines[0])
        if len(lines) > 1:
            contact_line("  |  ".join(lines[1:]))

        # ── SUMMARY ─────────────────────────────────────────────────────
        if sections.get("summary"):
            heading("Professional Summary")
            body_text(sections["summary"].strip())

        # ── CORE COMPETENCIES ───────────────────────────────────────────
        if sections.get("core_competencies"):
            heading("Core Competencies")
            body_text(sections["core_competencies"].strip(), color=GRAY)

        # ── EXPERIENCE ──────────────────────────────────────────────────
        exp_text = sections.get("experience", "")
        if exp_text:
            heading("Experience")
            current_job_bullets = []
            current_job_header  = None

            for line in exp_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # Detect job header lines: contain | or contain dates like "2020" or "–"
                is_job_header = (
                    "|" in line and ("20" in line or "–" in line or "-" in line)
                ) or (
                    ("–" in line or " - " in line) and
                    any(str(y) in line for y in range(2010, 2030)) and
                    len(line) < 120
                )
                if is_job_header:
                    # Flush previous job
                    if current_job_header:
                        parts = current_job_header.split("|")
                        if len(parts) >= 3:
                            job_title_line(parts[0].strip(), parts[1].strip(), parts[2].strip())
                        elif len(parts) == 2:
                            job_title_line(parts[0].strip(), parts[1].strip(), "")
                        else:
                            body_text(current_job_header, size=9.5)
                        for b in current_job_bullets:
                            bullet_line(b)
                    current_job_header  = line
                    current_job_bullets = []
                elif line.startswith("•") or line.startswith("-") or line.startswith("*"):
                    current_job_bullets.append(line.lstrip("•-* ").strip())
                else:
                    current_job_bullets.append(line)

            # Flush last job
            if current_job_header:
                parts = current_job_header.split("|")
                if len(parts) >= 3:
                    job_title_line(parts[0].strip(), parts[1].strip(), parts[2].strip())
                elif len(parts) == 2:
                    job_title_line(parts[0].strip(), parts[1].strip(), "")
                else:
                    body_text(current_job_header, size=9.5)
                for b in current_job_bullets:
                    bullet_line(b)

        # ── EDUCATION ───────────────────────────────────────────────────
        if sections.get("education"):
            heading("Education")
            for line in sections["education"].split("\n"):
                line = line.strip()
                if not line:
                    continue
                # Try to parse "Degree, Institution | Date" or plain
                if "|" in line:
                    parts = line.split("|")
                    deg_inst = parts[0].strip()
                    date     = parts[1].strip() if len(parts) > 1 else ""
                    if "," in deg_inst:
                        deg, inst = deg_inst.split(",", 1)
                        education_line(deg.strip(), inst.strip(), date)
                    else:
                        education_line(deg_inst, "", date)
                else:
                    body_text(line, size=9.5)

        # ── SKILLS ──────────────────────────────────────────────────────
        if sections.get("skills"):
            heading("Skills")
            for line in sections["skills"].split("\n"):
                line = line.strip()
                if line:
                    body_text(line, size=9.5, color=GRAY)

        # ── REMAINING SECTIONS ──────────────────────────────────────────
        done = {"header", "summary", "core_competencies", "experience", "education",
                "skills", "sections_changed", "changes_summary"}
        for key in ["projects", "certifications", "awards"]:
            if sections.get(key):
                heading(key.replace("_", " ").title())
                body_text(sections[key].strip(), size=9.5, color=GRAY)
                done.add(key)
        for key, val in sections.items():
            if key not in done and val and isinstance(val, str) and val.strip():
                heading(key.replace("_", " ").title())
                body_text(val.strip(), size=9.5)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _build_pdf(self, sections: dict, raw_text: str) -> bytes:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        VIOLET = colors.HexColor("#5B4FCF")
        BLACK  = colors.HexColor("#1A1A1A")
        GRAY   = colors.HexColor("#555555")

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=letter,
            leftMargin=0.75*inch, rightMargin=0.75*inch,
            topMargin=0.65*inch, bottomMargin=0.65*inch,
        )

        styles = getSampleStyleSheet()

        name_style = ParagraphStyle("Name", parent=styles["Normal"],
            fontSize=20, fontName="Helvetica-Bold", textColor=BLACK,
            alignment=TA_CENTER, spaceAfter=4)
        contact_style = ParagraphStyle("Contact", parent=styles["Normal"],
            fontSize=9, fontName="Helvetica", textColor=GRAY,
            alignment=TA_CENTER, spaceAfter=8)
        section_style = ParagraphStyle("Section", parent=styles["Normal"],
            fontSize=10, fontName="Helvetica-Bold", textColor=VIOLET,
            spaceBefore=10, spaceAfter=2)
        job_style = ParagraphStyle("Job", parent=styles["Normal"],
            fontSize=10, fontName="Helvetica-Bold", textColor=BLACK,
            spaceBefore=6, spaceAfter=2)
        body_style = ParagraphStyle("Body", parent=styles["Normal"],
            fontSize=9.5, fontName="Helvetica", textColor=BLACK,
            leading=13, spaceAfter=2)
        bullet_style = ParagraphStyle("Bullet", parent=styles["Normal"],
            fontSize=9.5, fontName="Helvetica", textColor=BLACK,
            leading=13, spaceAfter=2, leftIndent=12, firstLineIndent=-8)
        gray_style = ParagraphStyle("Gray", parent=styles["Normal"],
            fontSize=9, fontName="Helvetica", textColor=GRAY,
            leading=13, spaceAfter=2)

        def e(text):
            return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

        def add_heading(title):
            story.append(Paragraph(title.upper(), section_style))
            story.append(HRFlowable(width="100%", thickness=1, color=VIOLET, spaceAfter=3))

        story = []

        if not sections or len(sections) < 2:
            for line in raw_text.split("\n"):
                line = line.strip()
                if line:
                    story.append(Paragraph(e(line), body_style))
            doc.build(story)
            return buf.getvalue()

        # Header
        hdr_lines = [l.strip() for l in sections.get("header", "").split("\n") if l.strip()]
        if hdr_lines:
            story.append(Paragraph(e(hdr_lines[0]), name_style))
        if len(hdr_lines) > 1:
            story.append(Paragraph(e("  |  ".join(hdr_lines[1:])), contact_style))

        # Summary
        if sections.get("summary"):
            add_heading("Professional Summary")
            story.append(Paragraph(e(sections["summary"].strip()), body_style))
            story.append(Spacer(1, 4))

        # Core competencies
        if sections.get("core_competencies"):
            add_heading("Core Competencies")
            story.append(Paragraph(e(sections["core_competencies"].strip()), gray_style))
            story.append(Spacer(1, 4))

        # Experience
        exp_text = sections.get("experience", "")
        if exp_text:
            add_heading("Experience")
            for line in exp_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                is_job_header = (
                    "|" in line and ("20" in line or "–" in line)
                ) or (
                    ("–" in line or " - " in line) and
                    any(str(y) in line for y in range(2010, 2030)) and len(line) < 120
                )
                if is_job_header:
                    story.append(Paragraph(e(line), job_style))
                elif line.startswith("•") or line.startswith("-") or line.startswith("*"):
                    story.append(Paragraph("• " + e(line.lstrip("•-* ").strip()), bullet_style))
                else:
                    story.append(Paragraph("• " + e(line), bullet_style))
            story.append(Spacer(1, 4))

        # Education
        if sections.get("education"):
            add_heading("Education")
            for line in sections["education"].split("\n"):
                line = line.strip()
                if line:
                    story.append(Paragraph(e(line), body_style))
            story.append(Spacer(1, 4))

        # Skills
        if sections.get("skills"):
            add_heading("Skills")
            for line in sections["skills"].split("\n"):
                line = line.strip()
                if line:
                    story.append(Paragraph(e(line), gray_style))
            story.append(Spacer(1, 4))

        # Remaining
        done = {"header", "summary", "core_competencies", "experience", "education",
                "skills", "sections_changed", "changes_summary"}
        for key in ["projects", "certifications", "awards"]:
            if sections.get(key):
                add_heading(key.replace("_", " ").title())
                story.append(Paragraph(e(sections[key].strip()), gray_style))
                story.append(Spacer(1, 4))
                done.add(key)
        for key, val in sections.items():
            if key not in done and val and isinstance(val, str) and val.strip():
                add_heading(key.replace("_", " ").title())
                story.append(Paragraph(e(val.strip()), body_style))

        doc.build(story)
        return buf.getvalue()


resume_rewriter = ResumeRewriter()
