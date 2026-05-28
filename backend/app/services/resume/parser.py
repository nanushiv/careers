"""
Resume Parser — extracts structured text from PDF and DOCX files.
Uses PyMuPDF for PDF, python-docx for DOCX.
"""
import re
import logging
from pathlib import Path
from typing import Optional
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


class ParsedResume:
    def __init__(self):
        self.raw_text: str = ""
        self.sections: dict = {}
        self.word_count: int = 0
        self.error: Optional[str] = None

    def to_dict(self) -> dict:
        return {
            "raw_text": self.raw_text,
            "sections": self.sections,
            "word_count": self.word_count,
        }


SECTION_PATTERNS = {
    "experience": [
        r"(?i)\b(work\s+experience|professional\s+experience|employment|experience)\b",
    ],
    "education": [
        r"(?i)\b(education|academic\s+background|qualifications)\b",
    ],
    "skills": [
        r"(?i)\b(skills|technical\s+skills|core\s+competencies|technologies)\b",
    ],
    "summary": [
        r"(?i)\b(summary|profile|about\s+me|objective|overview)\b",
    ],
    "projects": [
        r"(?i)\b(projects|side\s+projects|notable\s+projects)\b",
    ],
    "certifications": [
        r"(?i)\b(certifications|certificates|licenses)\b",
    ],
}


class ResumeParser:

    def parse(self, file_path: str, file_type: str) -> ParsedResume:
        result = ParsedResume()
        try:
            if file_type == "pdf":
                result.raw_text = self._parse_pdf(file_path)
            elif file_type in ("docx", "doc"):
                result.raw_text = self._parse_docx(file_path)
            else:
                result.error = f"Unsupported file type: {file_type}"
                return result

            result.raw_text = self._clean_text(result.raw_text)
            result.word_count = len(result.raw_text.split())
            result.sections = self._extract_sections(result.raw_text)

        except Exception as e:
            logger.error(f"Resume parse error: {e}")
            result.error = str(e)

        return result

    def _parse_pdf(self, file_path: str) -> str:
        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        return "\n".join(text_parts)

    def _parse_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    def _clean_text(self, text: str) -> str:
        """Normalize whitespace and remove garbage characters."""
        # Normalize unicode
        text = text.encode("ascii", "ignore").decode("ascii")
        # Collapse whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        # Remove page numbers
        text = re.sub(r"^\d+\s*$", "", text, flags=re.MULTILINE)
        return text.strip()

    def _extract_sections(self, text: str) -> dict:
        """Split resume text into named sections."""
        sections = {}
        lines = text.split("\n")
        current_section = "header"
        current_content = []

        for line in lines:
            matched_section = None
            for section_name, patterns in SECTION_PATTERNS.items():
                for pattern in patterns:
                    if re.search(pattern, line) and len(line.strip()) < 50:
                        matched_section = section_name
                        break
                if matched_section:
                    break

            if matched_section:
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = matched_section
                current_content = []
            else:
                current_content.append(line)

        if current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections


resume_parser = ResumeParser()
